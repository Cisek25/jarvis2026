#!/usr/bin/env python3
"""Generate BRIEF_KLIENTA.docx — client questionnaire for IdoBooking websites."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

OUTPUT = Path(__file__).parent / "BRIEF_KLIENTA.docx"

ACCENT = RGBColor(0xA6, 0x7C, 0x52)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, color_hex: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    if num:
        run = p.add_run(f"  {num}  ")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
    run = p.add_run(f"  {text}")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT


def add_field(doc, label, hint=None, lines=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

    if hint:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        run2 = p2.add_run(hint)
        run2.font.italic = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = GRAY

    for _ in range(lines):
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(2)
        p3.paragraph_format.space_after = Pt(2)
        run3 = p3.add_run("_" * 80)
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def add_alert(doc, title, text, critical=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "F8D7DA" if critical else "FFF3CD")

    p = cell.paragraphs[0]
    run = p.add_run(f"{title}\n")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = DARK


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"\u2610  {item}")
        run.font.size = Pt(10)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("IDOBOOKING WEBSITE BUILDER")
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Brief Klienta")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Wypelnij formularz \u2014 na jego podstawie przygotujemy kompletna strone internetowa.")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    # ── 1. NAZWA I OPIS ──
    add_heading_styled(doc, "Nazwa i opis obiektu", "1")

    add_alert(doc,
        "Dane kontaktowe (adres, telefon, email, NIP)",
        "Te dane nalezy uzupelnic W PANELU IdoBooking, skad trafiaja automatycznie "
        "na strone (kontakt, stopka):\n\n"
        "Panel \u2192 Oferta \u2192 Miejsca noclegowe \u2192 lewy gorny rog: Lokalizacje "
        "\u2192 kliknij pierwsza lokalizacje \u2192 uzupelnij wszystkie pola\n\n"
        "Szczegolnie wazne: adres, telefon, email, wspolrzedne GPS."
    )

    add_field(doc, "Nazwa brandu / obiektu *", 'Jak ma sie nazywac strona (np. "Mountain Prestige")')
    add_field(doc, "Krotki opis obiektu (2-3 zdania) *", "Pojawi sie na stronie glownej w sekcji 'O nas'", lines=3)
    add_field(doc, "Co wyroznia obiekt? (3-6 kluczowych cech) *", "np. widok na gory, jacuzzi, blisko stoku, darmowy parking", lines=4)

    # ── 2. PANEL IDOBOOKING ──
    add_heading_styled(doc, "Panel IdoBooking \u2014 Miejsca noclegowe", "2")

    add_alert(doc,
        "KRYTYCZNE \u2014 Miejsca noclegowe musza byc dodane w panelu!",
        "Zanim zaczniemy budowac strone, nalezy dodac WSZYSTKIE apartamenty/pokoje "
        "w panelu IdoBooking:\n\n"
        "Panel \u2192 Oferta \u2192 Miejsca noclegowe \u2192 Dodaj miejsce noclegowe\n\n"
        "Kazdy apartament musi miec: nazwe, opis (min. 3-4 zdania), zdjecia (min. 5), "
        "ceny bazowe, max. liczbe gosci, udogodnienia z listy.\n\n"
        "BEZ TEGO STRONA BEDZIE PUSTA \u2014 nie bedzie ofert do wyswietlenia!",
        critical=True
    )

    add_field(doc, "ID klienta IdoBooking *", 'Widoczny w URL panelu, np. "57060"')
    add_field(doc, "Ile apartamentow/pokojow? *")
    add_field(doc, "Lista apartamentow (nazwy i max. osoby)", None, lines=5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Checklist \u2014 przed rozpoczeciem pracy:")
    run.font.bold = True
    run.font.size = Pt(10)

    add_checklist(doc, [
        "Dodalem wszystkie miejsca noclegowe w panelu (Oferta \u2192 Miejsca noclegowe)",
        "Kazde miejsce ma min. 5 zdjec i opis",
        "Ustawilem ceny bazowe",
        "Zaznaczylem udogodnienia (Wi-Fi, parking, TV itd.)",
        "Wyroznilem oferty do pokazania na glownej (Oferta \u2192 Wyroznij)",
        "Uzupelnilem dane lokalizacji (Miejsca noclegowe \u2192 Lokalizacje \u2192 szczegoly)",
    ])

    # ── 3. DOMENA ──
    add_heading_styled(doc, "Domena i jezyki", "3")
    add_field(doc, "Domena (jesli posiadasz)", "np. mountainprestige.pl")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Jezyki strony:")
    run.font.bold = True
    run.font.size = Pt(10)
    add_checklist(doc, ["Polski (PL)", "Angielski (EN)", "Niemiecki (DE)", "Inny: ___________"])

    # ── 4. DESIGN ──
    add_heading_styled(doc, "Styl wizualny", "4")

    p = doc.add_paragraph()
    run = p.add_run("Wybierz gotowa palete kolorow lub opisz preferencje:")
    run.font.size = Pt(10)

    add_checklist(doc, [
        "A. Warm Copper \u2014 ciepla, premium, gory (zloto-miedz + krem)",
        "B. Forest Green \u2014 natura, eco, camping (zielen lesna)",
        "C. Elegant Gold \u2014 luksus, city, modern (zloto + biel)",
        "D. Coastal Blue \u2014 morze, jeziorka (blekit + granat)",
        "E. Warm Terracotta \u2014 rustykalny, wiejski (terakota + krem)",
        "Wlasna paleta (opisz ponizej)",
    ])

    add_field(doc, "Wlasne kolory (opcjonalnie)",
              "Mozesz podac kody hex (np. #A67C52) albo opisac slownie (np. 'cieply zloty', 'zielen lesna')")

    add_field(doc, "Strony / marki ktore Ci sie podobaja (inspiracje)", "Linki do stron, screenshoty", lines=2)

    # ── 5. ZDJECIA ──
    add_heading_styled(doc, "Zdjecia i logo", "5")

    add_alert(doc,
        "Zdjecia to 80% sukcesu strony!",
        "Prosimy o przeslanie zdjec w NAJWYZSZEJ dostepnej rozdzielczosci "
        "(min. 1200px szerokosc). Najlepiej jako link do Google Drive / WeTransfer."
    )

    table = doc.add_table(rows=7, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Co", "Ile", "Uwagi", "Mam?"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "F5F3EF")
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    data = [
        ("Logo", "1", "PNG/SVG z przezroczystym tlem, min. 400px", "\u2610"),
        ("Slider (hero)", "2-4", "Panoramiczne 1920x1080+", "\u2610"),
        ("O nas", "1-2", "Budynek, lokalizacja", "\u2610"),
        ("Apartamenty", "5-10/apt.", "Juz dodane w panelu IdoBooking", "\u2610"),
        ("Lokalizacja", "3-5", "Okolica, atrakcje, mapa", "\u2610"),
        ("Galeria", "6-12", "Mix: wnetrza + okolica + detale", "\u2610"),
    ]
    for r, (co, ile, uwagi, mam) in enumerate(data, 1):
        table.rows[r].cells[0].text = co
        table.rows[r].cells[1].text = ile
        table.rows[r].cells[2].text = uwagi
        table.rows[r].cells[3].text = mam
        for c in range(4):
            for par in table.rows[r].cells[c].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)

    add_field(doc, "Link do zdjec (Google Drive / WeTransfer) *", None)

    # ── 6. TRESCI ──
    add_heading_styled(doc, "Tresci i podstrony", "6")

    p = doc.add_paragraph()
    run = p.add_run("Podstrony na stronie (zaznacz potrzebne):")
    run.font.bold = True
    run.font.size = Pt(10)

    add_checklist(doc, [
        "Oferta (lista apartamentow \u2014 automatyczna)",
        "Kontakt (formularz \u2014 automatyczny)",
        "Lokalizacja (mapa, atrakcje, dojazd)",
        "Wspolpraca (dla wlascicieli apartamentow)",
        "Galeria",
        "Blog",
        "FAQ (najczesciej zadawane pytania)",
        "Inna: ___________",
    ])

    add_field(doc, "Tresc FAQ (pytania i odpowiedzi)", None, lines=4)
    add_field(doc, "Tresc 'O lokalizacji'", "Gdzie jest obiekt, co w poblizu, odleglosci", lines=3)
    add_field(doc, "Tresc 'Wspolpraca'", "Informacja dla wlascicieli apartamentow", lines=3)

    # ── 7. SOCIAL MEDIA ──
    add_heading_styled(doc, "Social media i integracje", "7")
    add_field(doc, "Facebook")
    add_field(doc, "Instagram")
    add_field(doc, "Google Maps (link do obiektu)")
    add_field(doc, "Booking.com (link do obiektu)")

    # ── 8. UWAGI ──
    add_heading_styled(doc, "Uwagi dodatkowe", "8")
    add_field(doc, "Dodatkowe informacje, zyczenia, uwagi", None, lines=5)
    add_field(doc, "Termin \u2014 kiedy strona powinna byc gotowa?")

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    run = p.add_run("Wypelniony brief prosimy wyslac na adres: [TWOJ EMAIL]")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.save(str(OUTPUT))
    print(f"  OK  {OUTPUT.name} saved ({OUTPUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
